import os
import glob
import json
# import logging
import pandas as pd
from docx import Document
from docx.shared import Inches
from loguru import logger
# https://cloud.tencent.com/developer/article/1763991
from xpinyin import Pinyin
from utils import prompts, gpt3, stable_diffusion, doc, sqlite
from gen_sd_prompts import text_generate
"""
pip install python-docx
todo 仍有机器人的不恰当的回答未过滤, 详见 [元]潘纯_2023-12-15.docx
todo 单句生成的图片和整体理解的意思不对应
2023-12-16 由于图像生成过于随意, 将逐句生成改为整首诗生成
todo 汉、唐、宋、明、魏晋 可以尝试使用 lora hanfu_v30 (发现在prompt中增加lora标签无效)
2023-12-22 增加进度记录和断点继续
"""
# logging.root.setLevel(logging.NOTSET)
# logging.basicConfig(format="%(thread)d %(asctime)s %(name)s:%(levelname)s:%(lineno)d:%(message)s",
#                     datefmt="%Y-%m-%d %H:%M:%S", level=logging.DEBUG)
# logger = logging.getLogger(__name__)
logger.add("logs/main.log", rotation="50 MB")
# todo 使用了lora后, 效果反而变差, 检查原因
lora_styles = {"唐": ", tang style, <lora:hanfu_v30:1>"}
has_gen_author = False
progress_record = {"book_idx": 0, "author_idx": 0, "poem_idx": 0, "part_idx": 0}


def load_progress():
    """
    加载进度
    """
    global progress_record
    if os.path.exists("progress.json"):
        with open("progress.json", "r", encoding="utf-8") as f:
            progress_record = json.load(f)
    return progress_record


def save_progress():
    """
    保存进度
    """
    global progress_record
    with open("progress.json", "w", encoding="utf-8") as f:
        json.dump(progress_record, f, ensure_ascii=False, indent=4)


def read_data():
    """
    将唐诗宋词读入
    """
    input_dir = "input/"
    books = glob.glob(input_dir + "*.csv")
    books.sort()
    return books


def get_llm_lines(content, delimiters = ["；", "。"]):
    # 2023-12-15发现LLM对单句理解不到位, 改为一句话
    string = content
    for delimiter in delimiters:
        string = ";".join(string.split(delimiter))
    result = string.split(";")
    result = [r.strip() for r in result if len(r.strip())>0]
    return result


def gen_image(en_words):
    image_url = None
    sd_prompt = text_generate(en_words)
    try:
        print("use prompt to gen pic: ", sd_prompt)
        # 这里无需在prompts中设置风格, 减轻模型计算负担, 直接在sd webui api中指定lora
        # image_prompt = prompts.illustration(f"{plot}\n\n{'' if i == 0 else parts[i - 1]}\n\n{part}")
        image_url = stable_diffusion.generate_image(sd_prompt)
    except Exception as e:
        print(e)
    return image_url


def get_author_column(se, col_name, lang_type="cn"):
    if lang_type=="cn":
        col_value = '不详'
    else:
        col_value = 'Unknown'
    if se[col_name] is not None:
        col_value = str(se[col_name])
    return col_value


def add_author_bio(document, se_author):
    """
    组织成可展示的格式
    """
    ret_sex = int(se_author['c_female']) 
    author_sex = "男" if ret_sex == 0 else "女"
    author_sex_en = "Male" if ret_sex == 0 else "Female"
    document.add_paragraph("性别: " + author_sex)
    document.add_paragraph("Gender: " + author_sex_en)
    birth_year = get_author_column(se_author, 'c_birthyear')
    birth_year_en = get_author_column(se_author, 'c_birthyear', 'en')
    death_year = get_author_column(se_author, 'c_deathyear')
    death_year_en = get_author_column(se_author, 'c_deathyear', 'en')
    document.add_paragraph("生于: " + birth_year + " 卒于: " + death_year)
    document.add_paragraph("Born: " + birth_year_en + " Died: " + death_year_en)
    life_story = get_author_column(se_author, 'c_fl_ey_notes')
    life_story_en = ""
    if len(life_story)>0 and life_story!='不详':
        try:
            life_prompt = prompts.title_translation(life_story)
            life_story_en = gpt3.generate_with_prompt(life_prompt, 0.6)
            # 如果仍然包含多余的内容, 则手工去除
            if life_story_en.find('"')>0:
                life_story_en = life_story_en.split('"')[1].strip()
            print("translate life story:", life_story_en)
        except Exception as e:
            logger.exception(e)
    elif life_story=='不详':
        life_story_en = "Unknown"
    document.add_paragraph("生平事迹: " + life_story)
    document.add_paragraph("Life story: " + life_story_en)
    return ret_sex


def get_first_result(llm_result, text_in_symbol='"', default_value=None):
    result = default_value
    if llm_result.find(text_in_symbol)>0:
        result = llm_result.split('"')[1].strip()
    else:
        lines = get_llm_lines(llm_result, [".", "。", "#"])
        result = lines[0]
    return result


def gen_author(se, document, title, content):
    """
    生成作者信息(单独做成1页)
    """
    global lora_styles
    chinese_dynasty = str(se['朝代'])
    try:
        dynasty_prompt = prompts.dynasty_translation(chinese_dynasty)
        dynasty = gpt3.generate_with_prompt(dynasty_prompt, 0.6)
        dynasty = get_first_result(dynasty, text_in_symbol='"', default_value=chinese_dynasty)
        if dynasty==chinese_dynasty:
            p = Pinyin()
            dynasty = p.get_pinyin(chinese_dynasty)
        logger.debug("translate dynasty: " + dynasty)
    except Exception as e:
        logger.exception(e)

    chinese_author = str(se['作者'])
    try:
        author_prompt = prompts.author_translation(chinese_author)
        author = gpt3.generate_with_prompt(author_prompt, 0.6)
        author = get_first_result(author, text_in_symbol='"', default_value=chinese_author)
        if author==chinese_author:
            p = Pinyin()
            author = p.get_pinyin(chinese_author, tone_marks='marks')
        logger.debug("translate author: " + author)
    except Exception as e:
        logger.exception(e)

    document.add_heading(chinese_author, 0)
    document.add_heading(author, 1)

    se_author = sqlite.get_author_info(chinese_author, chinese_dynasty)
    try:
        gender = 'male' if int(se_author['c_female'])==0 else 'female'
    except Exception:
        gender = 'male'
    
    author_image = None
    try:
        portrait_prompt = prompts.author_profile(title, dynasty, author, gender, content)
        portrait = gpt3.generate_with_prompt(portrait_prompt, 0.6)
        portrait = get_first_result(portrait, text_in_symbol='"', default_value=portrait)
        # 增加lora调用
        portrait += lora_styles[chinese_dynasty]
        logger.debug("author profile: " + portrait)
        author_image = gen_image(portrait)
    except Exception as e:
        logger.exception(e)
    
    table = document.add_table(rows=1, cols=2)
    head_row = table.rows[0].cells
    if author_image is not None:
        try:
            paragraph = head_row[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(author_image, width=Inches(3), height=Inches(3))
            # poem_pics.append(author_image)
            doc.set_cell_margins(head_row[0], top=0, start=0, bottom=0, end=0)
        except Exception as e:
            logger.exception(e)
    head_row[1].paragraphs[0].text = '朝代: ' + chinese_dynasty
    head_row[1].add_paragraph('Dynasty: ' + dynasty)
    if se_author is not None:
        add_author_bio(head_row[1], se_author)

    document.add_page_break()
    return author_image


def gen_story(se, document, author=None, doc_name=None):
    global has_gen_author, progress_record
    chinese_title = str(se['题目'])
    try:
        title_prompt = prompts.title_translation(chinese_title)
        title = gpt3.generate_with_prompt(title_prompt, 0.6)
        title = get_first_result(title, text_in_symbol='"', default_value=title)
        logger.debug("translate title: " + title)
    except Exception as e:
        logger.exception(e)
        return False

    chinese_content = str(se['内容'])
    try:
        content_prompt = prompts.poem_translation(chinese_content)
        content = gpt3.generate_with_prompt(content_prompt, 0.6)
        content = get_first_result(content, text_in_symbol='"', default_value=content)
        logger.debug("translate content: " + content)
    except Exception as e:
        logger.exception(e)
    
    poem_pics = []
    # 如果是某位作者的专辑, 则只需要生成1次作者简介
    if not has_gen_author:
        author_image = gen_author(se, document, title, content)
        if author_image is not None:
            poem_pics = [author_image]
        if author is not None:
            has_gen_author = True

    if doc_name is not None and progress_record['part_idx']<1:
        document.save(doc_name)
        progress_record['part_idx'] = 1
        save_progress()

    document.add_heading(chinese_title, 1)
    document.add_heading(title, 2)

    table = document.add_table(rows=1, cols=2)

    try:
        poem_view_prompt = prompts.poem_view(content)
        poem_view = gpt3.generate_with_prompt(poem_view_prompt, 0.6)
        poem_view = get_first_result(poem_view, text_in_symbol='"', default_value=poem_view)
        logger.debug("poem view: " + poem_view)
    except Exception as e:
        logger.exception(e)

    image_url = gen_image(poem_view)
    if image_url is not None:
        try:
            paragraph = table.rows[0].cells[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_url, width=Inches(3), height=Inches(3))
            poem_pics.append(image_url)
            doc.set_cell_margins(table.rows[0].cells[0], top=0, start=0, bottom=0, end=0)
        except Exception as e:
            logger.exception(e)

    table.rows[0].cells[1].paragraphs[0].text = chinese_content
    table.rows[0].cells[1].add_paragraph(content)

    if doc_name is not None and progress_record['part_idx']<2:
        document.save(doc_name)
        progress_record['part_idx'] = 2
        save_progress()
    for p in poem_pics:
        try:
            os.remove(p)
        except Exception:
            pass
    # 解释整首诗的意思
    document.add_heading("诗句解析(Poetry analysis)", 3)
    try:
        interpret_prompt = prompts.poem_interpret(chinese_content)
        interpret = gpt3.generate_with_prompt(interpret_prompt, 0.6)
        logger.debug("interpret content: " + interpret)
    except Exception as e:
        logger.exception(e)
    document.add_paragraph(interpret)
    # 想象一个小故事
    document.add_heading("诗句背后的故事(The story behind the poem)", 3)
    try:
        story_prompt = prompts.poem_associate_story(chinese_content)
        story = gpt3.generate_with_prompt(story_prompt, 0.6)
        # 故事不能用""过滤
        logger.debug("associate story: " + story)
    except Exception as e:
        logger.exception(e)
    document.add_paragraph(story)
    document.add_page_break()
    if doc_name is not None and progress_record['part_idx']<3:
        document.save(doc_name)
        progress_record['part_idx'] = 3
        save_progress()


def gen_one_book(df_story, book, author):
    """
    生成一本书, 生成根据作者分成多本
    """
    global progress_record
    book_name = book.split("/")[-1].split("\\")[-1].split(".")[0]
    doc_name = 'output/[' + book_name + ']' + author + '.docx'
    document = Document()
    # 为了谨慎使用接口, 以及方便对author信息输出做控制, 这里使用for循环
    for index, row in df_story.iterrows():
        if index < progress_record["poem_idx"]:
            continue
        gen_story(row, document, author, doc_name)
        progress_record["poem_idx"] = progress_record["poem_idx"] + 1
        progress_record["part_idx"] = 0
        save_progress()
        # df_story.apply(gen_story, axis=1, document=document, author=author)
    # document.save(doc_name)
    progress_record['author_idx'] = progress_record['author_idx'] + 1
    progress_record['poem_idx'] = 0
    progress_record['part_idx'] = 0
    save_progress()
    logger.debug("book " + doc_name + " generated")


if __name__ == "__main__":
    load_progress()
    # books = read_data()
    book = "input/唐.csv"
    df_story = pd.read_csv(book)
    author = "李白"
    df_author = df_story[df_story['作者']==author]
    gen_one_book(df_author, book, author)
    progress_record['book_idx'] = progress_record['book_idx'] + 1
    progress_record['author_idx'] = 0
    progress_record["poem_idx"] = 0
    progress_record['part_idx'] = 0
    save_progress()
    # for b_index, book in enumerate(books):
        # df_story = pd.read_csv(book)
        # authors = df_story['作者'].unique().tolist()
        # authors.sort()
        # 2023-12-13 由于目测某些作者写作内容较多, 大于10篇的按作者划分, 不够的作者综合成1篇
        # list_other = []
        # for i, author in enumerate(authors):
        #     has_gen_author = False
        #     df_author = df_story[df_story['作者'] == author]
        #     if len(df_author) >= 10:
        #         gen_one_book(df_author, book, author)
        #         logger.debug("author(" + str(i) + ") " + author + " generated")
        #     else:
        #         list_other.append(df_author)
        # if len(list_other) > 0:
        #     df_other = pd.concat(list_other)
        #     has_gen_author = False
        #     gen_one_book(df_other, book, "其他")
        # logger.debug("book(" + str(b_index) + ") " + book + " generated")
